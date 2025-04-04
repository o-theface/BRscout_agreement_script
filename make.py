# -----------------------------------------------------------------------------
# Copyright (c) 2025, Jean Carlo dos Santos (o_theface@yahoo.com.br)
# All rights reserved.
#
# This software is licensed under the terms of the MIT License.
# For a copy, see the LICENSE file in the root directory of this project.
#
# Description: Script to generate personalized agreements based on a template
# and data from a CSV file.
#
# This code use this library: pandas e python-docx.
# To install it, run:
# pip install pandas python-docx
# -----------------------------------------------------------------------------

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

import os

# Load associates data from CSV file
data = pd.read_csv("dados_associados.csv", sep=";")

# load the template document
doc_model = Document("acordo_mutuo.docx")

# Datas fixas
START_DATE = "01/01/2025"
END_DATE = "31/01/2026"
SIGN_DATE = "01/04/2025"

# make sure the output directory exists
output_dir = "acordos_gerados"
os.makedirs(output_dir, exist_ok=True)

# Function to generate agreements
def gerar_acordos(data, doc_model):
    for _, row in data.iterrows():
        name = row["Nome do associado"]
        adress = row["Endereço"]
        ident_id = row["CPF"]

        for i in range(1, 4):  # Assuming there are 3 functions to process
            category = row.get(f"Categoria - {i}º Função", "")
            section = row.get(f"Ramo - {i}º Função", "")
            function = row.get(f"Função - {i}º Função", "")

            if ((category == "Não se aplica") or (function == "funcao")):
                continue
            if (section != "Não se aplica"):
                    function += f" - {section}"

            if pd.notna(category) and pd.notna(section) and pd.notna(function):
                new_doc = Document()
                section = new_doc.sections[0]
                # Define the page size (A4)
                section.page_width = Inches(8.27)
                section.page_height = Inches(11.69)
                # Set margins (in inches)
                section.top_margin = Inches(1)       # 1 inch top margin
                section.bottom_margin = Inches(1)    # 1 inch bottom margin
                section.left_margin = Inches(1)   # 1.25 inches left margin
                section.right_margin = Inches(1) 
                isFirstParagraph = True
                for p in doc_model.paragraphs:
                    text = p.text
                    text = text.replace("{Nome do associado}", name)
                    text = text.replace("{Endereco}", adress)
                    text = text.replace("{CPF}", str(ident_id))
                    text = text.replace("{Categoria}", category)
                    text = text.replace("{Função}", function)
                    text = text.replace("{DATA_INI}", START_DATE)
                    text = text.replace("{DATA_FIM}", END_DATE)
                    text = text.replace("{DATA_ASS}", SIGN_DATE)
                    
                    new_paragraph = new_doc.add_paragraph()
                    run = new_paragraph.add_run(text)
                    run.font.name = "Arial"  # Set font family
                    
                    # Add bold to "CLÁUSULA" paragraph
                    if ((text.find("CLÁUSULA") == 0) or isFirstParagraph):
                        run.bold = True

                     # Add paragraph with controlled spacing
                    
                    if (isFirstParagraph == True):
                        new_paragraph.paragraph_format.space_after = Pt(12)
                        new_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(14)
                        isFirstParagraph = False
                    elif (text == ""):
                        run.font.size = Pt(5)
                        new_paragraph.paragraph_format.space_after = Pt(0) 
                    else:
                        new_paragraph.paragraph_format.space_after = Pt(0) 
                        new_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        run.font.size = Pt(12)

                # Saving new aggreement document
                safe_name = name.replace(" ", "_").lower()
                file_name = f"{output_dir}/acordo_{safe_name}_{i}.docx"
                new_doc.save(file_name)
                print(f"Acordo gerado: {file_name}")
    
gerar_acordos(data, doc_model)
print("Todos os acordos foram gerados com sucesso.")